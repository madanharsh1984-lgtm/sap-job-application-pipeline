"""
resume_builder.py — JD-Tailored ATS Resume Generator
======================================================
Generates a tailored, ATS-friendly DOCX resume for Harsh Madan
based on the recruiter's job post text.

How it works:
  1. Parse the JD text → detect role type, modules, keywords
  2. Select the best-matching summary + reorder bullet points
  3. Build a clean ATS DOCX (no tables, no text boxes, no images)
  4. Return the file path for attachment

ATS rules enforced:
  - Single-column layout (no tables for layout)
  - Standard section headings (no icons, no graphics)
  - All text in paragraph runs (no text boxes)
  - Keywords from JD injected into summary and skills
  - File saved as .docx (not .doc or .pdf)
  - Standard fonts: Calibri / Arial
"""

import os, re, io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r"C:\Users\madan\OneDrive\Desktop\Linkdin Job Application"
TEMP_DIR = os.path.join(BASE_DIR, "temp_resumes")
os.makedirs(TEMP_DIR, exist_ok=True)

# ── CANDIDATE CONSTANTS ──────────────────────────────────────────────────────
NAME         = "HARSH MADAN"
TITLE_BASE   = "SAP S/4HANA Program Manager  |  Data Migration Specialist  |  15+ Years"
CONTACT      = "+91 96679 64756  •  Madan.harsh1984@gmail.com  •  linkedin.com/in/harsh-madan-b818113b  •  Delhi NCR / Remote"
AVAILABILITY = "Notice Period: Immediate  |  Current CTC: 35 LPA  |  Expected CTC: 40 LPA  |  Open to: Remote / Hybrid"

# ── ROLE DETECTION ───────────────────────────────────────────────────────────
ROLE_KEYWORDS = {
    "data_migration": [
        "data migration", "lsmw", "ltmc", "slt", "migration", "etl",
        "data transfer", "legacy data", "cutover", "data quality",
        "master data", "mdg", "data governance", "idoc", "bapi"
    ],
    "program_manager": [
        "program manager", "project manager", "pmo", "programme manager",
        "s/4hana", "s4hana", "transformation", "rollout", "implementation lead",
        "delivery manager", "engagement manager"
    ],
    "fico": [
        "fico", "fi/co", "fi-co", "financial accounting", "controlling",
        "gl", "general ledger", "ap ", "accounts payable", "ar ",
        "accounts receivable", "asset accounting", "aa ", "co-pa",
        "profit center", "cost center", "period close"
    ],
    "mm": [
        " mm ", "materials management", "procurement", "purchase order",
        "inventory", "warehouse", "mrp", "goods receipt", "fi-mm",
        "vendor management", "sourcing"
    ],
    "sd": [
        " sd ", "sales distribution", "order management", "billing",
        "sales order", "fi-sd", "pricing", "revenue recognition", "crm"
    ],
}

def detect_roles(jd_text: str) -> dict:
    """Return dict of role → score based on keyword hits in JD."""
    text = jd_text.lower()
    scores = {}
    for role, keywords in ROLE_KEYWORDS.items():
        scores[role] = sum(1 for kw in keywords if kw in text)
    return scores

def extract_jd_keywords(jd_text: str) -> list:
    """Extract SAP-specific keywords from JD for ATS injection."""
    SAP_TERMS = [
        "SAP S/4HANA", "SAP ECC", "S/4HANA", "FICO", "FI/CO",
        "MM", "SD", "MDG", "LSMW", "LTMC", "SLT", "SAP CPI",
        "SAP Activate", "SAP Solution Manager", "CHARM",
        "Data Migration", "Cutover", "SIT", "UAT", "PMO",
        "Go-Live", "BAPI", "IDOC", "ABAP", "Fiori",
        "Master Data", "SAP MDG", "BTP", "SAP BTP",
        "Agile", "JIRA", "Smartsheet", "Waterfall",
        "Stakeholder Management", "Change Management",
        "Business Blueprint", "Fit-Gap", "Hypercare",
        "Period Close", "Financial Close", "SAP GRC",
        "SAP APO", "SAP IBP", "SAP WM", "SAP EWM",
        "SAP HR", "SAP HCM", "SAP SuccessFactors",
    ]
    text_lower = jd_text.lower()
    found = []
    for term in SAP_TERMS:
        if term.lower() in text_lower and term not in found:
            found.append(term)
    return found

def detect_location_preference(jd_text: str) -> str:
    text = jd_text.lower()
    if "remote" in text:
        return "Remote / Hybrid"
    if "hybrid" in text:
        return "Hybrid / On-site"
    return "Remote / Hybrid"

# ── SUMMARY VARIANTS ─────────────────────────────────────────────────────────
SUMMARIES = {
    "data_migration": (
        "SAP Data Migration Specialist with 15+ years architecting and delivering enterprise-scale "
        "ECC → S/4HANA data migrations. Expert in LSMW, LTMC, SLT, SAP CPI, and BAPI-based ETL "
        "pipelines. Migrated 50M+ records across 30+ countries with 100% data reconciliation accuracy "
        "and 35% faster cycle time vs. industry benchmark. Deep expertise in master data governance "
        "(SAP MDG), cutover planning, reconciliation frameworks, and post-migration hypercare. "
        "Immediate joiner • Remote • Expected CTC: 40 LPA."
    ),
    "program_manager": (
        "SAP S/4HANA Program Manager with 15+ years leading global ERP transformations from blueprint "
        "through go-live. Built and managed 80+ consultant PMOs across FICO, MM, SD, and MDG workstreams. "
        "Delivered 12 global S/4HANA go-lives across 30 countries — on time, within budget, zero critical "
        "post-go-live defects. Expert in SAP Activate, cutover governance, SIT/UAT frameworks, and "
        "C-suite stakeholder management. Immediate joiner • Remote • Expected CTC: 40 LPA."
    ),
    "fico": (
        "SAP FICO Program Manager & Implementation Lead with 15+ years delivering end-to-end SAP FICO "
        "solutions across GL, AP, AR, AA, CO-PA, and Profit Centre Accounting. Led 18-branch pan-India "
        "FICO rollout (Sony India) and global S/4HANA FICO migration at Autodesk. Expert in FI-MM and "
        "FI-SD integration, period-end close automation, and SAP MDG for master data governance. "
        "Reduced close cycle from 12 to 5 days. Immediate joiner • Remote • Expected CTC: 40 LPA."
    ),
    "mm": (
        "SAP MM / Procurement Functional Lead with 15+ years delivering SAP MM implementations and "
        "S/4HANA transformations. Deep expertise in Procurement, Inventory Management, MRP, FI-MM "
        "integration, and vendor master data governance via SAP MDG. Led data migration programs "
        "(LSMW, LTMC) for material masters, vendor masters, and open POs across global rollouts. "
        "Immediate joiner • Remote • Expected CTC: 40 LPA."
    ),
    "sd": (
        "SAP SD / Order-to-Cash Specialist with 15+ years across sales order management, billing, "
        "pricing, FI-SD integration, and S/4HANA rollouts. Delivered cross-functional PMO spanning "
        "SD, FICO, and MM with full SIT/UAT governance and cutover planning. Expert in data migration "
        "of customer masters and open sales orders via LSMW and LTMC. "
        "Immediate joiner • Remote • Expected CTC: 40 LPA."
    ),
    "default": (
        "SAP S/4HANA Program Manager with 15+ years driving large-scale ERP transformations across "
        "Manufacturing, Technology, and Consumer Electronics sectors. Specialises in ECC → S/4HANA "
        "migrations, enterprise data migration programs (LSMW, LTMC, SLT, SAP CPI), global rollouts, "
        "and high-performance SAP PMOs. Delivered 12 global go-lives, migrated 50M+ records, generated "
        "₹4 Cr+ annual savings through process standardisation. "
        "Immediate joiner • Remote • Expected CTC: 40 LPA."
    ),
}

def pick_summary(scores: dict) -> str:
    """Pick best summary based on highest scoring role."""
    # Data migration always takes priority if strongly present
    if scores.get("data_migration", 0) >= 2:
        return SUMMARIES["data_migration"]
    top = max(scores, key=scores.get)
    if scores[top] == 0:
        return SUMMARIES["default"]
    return SUMMARIES.get(top, SUMMARIES["default"])

# ── SKILLS TABLE (ATS-friendly: listed as text, not a table) ─────────────────
SKILLS_SECTIONS = {
    "SAP Modules": "FICO (GL, AP, AR, AA, CO, CO-PA) | MM | SD | MDG | SAP CPI | SAP SLT",
    "Data Migration": "LSMW | LTMC | SAP SLT | SAP CPI | BAPI | IDOC | ETL | Data Reconciliation | Cutover Planning",
    "Programme Delivery": "SAP Activate | PMO Leadership | SIT / UAT | Defect Triage | Cutover Runbooks | Hypercare | Change Management",
    "Tools & Platforms": "SAP S/4HANA | SAP ECC 6.0 | SAP Solution Manager | CHARM | JIRA | Smartsheet | SAP Fiori",
    "Master Data": "SAP MDG | Vendor Master | Customer Master | Material Master | Chart of Accounts | Data Governance",
}

# ── KEY ACHIEVEMENTS ─────────────────────────────────────────────────────────
ALL_ACHIEVEMENTS = [
    {
        "roles": ["program_manager", "data_migration"],
        "text": "Zero-Defect Global Go-Live (Autodesk): Led cutover for 30-country S/4HANA rollout under a tight 6-month timeline. Built 1,200-script UAT framework and parallel-run governance. Result: 12 go-lives on time with ZERO critical post-go-live defects."
    },
    {
        "roles": ["data_migration"],
        "text": "50M+ Record Data Migration (Autodesk): Designed end-to-end migration architecture using LTMC, SLT, and SAP CPI for financial, master, and transactional data across 30+ countries. Result: 100% data reconciliation accuracy; 35% faster cycle time vs. industry benchmark."
    },
    {
        "roles": ["fico", "program_manager"],
        "text": "₹4 Cr Annual Cost Savings (Autodesk): Implemented SAP MDG to centralise master data and automated period-close via FICO configuration. Result: ₹4 Cr/year operational savings; 12 FTE reduction across finance BPO."
    },
    {
        "roles": ["program_manager", "fico", "data_migration"],
        "text": "97% UAT First-Pass Rate (Autodesk): Built SIT/UAT governance framework for 80+ consultants across FICO, MM, SD, MDG. Result: 97% first-pass rate across 1,200+ test scripts — highest in project history."
    },
    {
        "roles": ["fico", "mm", "sd"],
        "text": "18-Branch SAP FICO Rollout (Sony India): Led pan-India SAP FICO rollout (GL, AP, AR, AA, CO) across 18 branches. Migrated vendor/customer masters via LSMW. Result: On-time delivery, 500+ users trained, zero post-go-live escalations."
    },
    {
        "roles": ["mm"],
        "text": "FI-MM Integration & Inventory Rationalisation (iGATE Patni): Designed FI-MM blueprints and resolved 40+ cross-module integration gaps during SIT for 3 mid-market manufacturing clients."
    },
    {
        "roles": ["sd"],
        "text": "FI-SD Integration & Revenue Recognition (iGATE Patni): Delivered FI-SD integration blueprints and order-to-cash configuration for manufacturing clients; resolved 40+ cross-module gaps, preventing go-live delays."
    },
]

def rank_achievements(scores: dict) -> list:
    """Return achievements sorted by relevance to detected roles."""
    top_roles = sorted(scores, key=scores.get, reverse=True)

    def score_achievement(ach):
        return sum(scores.get(r, 0) for r in ach["roles"])

    ranked = sorted(ALL_ACHIEVEMENTS, key=score_achievement, reverse=True)
    # Always show top 5
    return [a["text"] for a in ranked[:5]]

# ── EXPERIENCE BULLETS ───────────────────────────────────────────────────────
EXP_AUTODESK = [
    "Architected enterprise data migration strategy using LTMC, SLT, and SAP CPI — migrated 50M+ financial, master, and transactional records with 100% reconciliation accuracy and 35% faster cycle time vs. industry benchmark.",
    "Built and led a PMO of 80+ SAP consultants; designed phased S/4HANA migration roadmap using SAP Activate methodology. All 12 go-lives delivered on schedule with zero critical post-go-live defects.",
    "Built SIT/UAT governance framework with 1,200+ test scripts across FICO, MM, SD, MDG workstreams; achieved 97% first-pass rate — highest in Autodesk project history.",
    "Implemented SAP MDG (Master Data Governance) to centralise vendor/customer/material masters; eliminated 12 FTEs of manual reconciliation, generating ₹4 Cr annual savings.",
    "Drove cutover planning for 12 global go-lives: developed cutover runbooks, managed parallel-run windows, led hypercare war rooms, ensured <2-hour RTO across all geographies.",
    "Managed SAP Solution Manager for CHARM, defect tracking, test management, and release management across the full programme lifecycle.",
    "Presented monthly programme status to C-suite (CFO, CIO, VP Finance); managed $15M+ programme budget with zero overrun.",
]

EXP_LAVA = [
    "Led chart of accounts restructuring and intercompany reconciliation for 8 entities; reduced month-end close from 12 days to 5 days using automated accrual and allocation cycles.",
    "Executed LSMW/BAPI data migration for vendor masters, customer masters, asset register, and open items — 2M+ records migrated with zero data quality escalations.",
    "Coordinated UAT with 60+ finance and operations stakeholders; resolved 280 defects in 6-week SIT cycle, enabling on-time go-live.",
]

EXP_SONY = [
    "Led pan-India SAP FICO rollout (GL, AP, AR, AA, CO) across 18 branch offices for 500+ end users.",
    "Migrated vendor/customer masters, open items, and fixed asset register using LSMW; achieved 100% data accuracy post-migration with zero rollback.",
    "Served as primary liaison between business finance leadership and SAP technical team; resolved 150+ configuration gaps during blueprint and realisation phases.",
    "Delivered SAP end-user training for 500+ finance users across 18 locations; post-go-live support satisfaction score: 4.7/5.",
]

EXP_IGATE = [
    "Designed FI-MM and FI-SD integration blueprints for 3 mid-market clients; resolved 40+ cross-module gaps during SIT, preventing go-live delays.",
    "Conducted business process workshops and blueprint documentation for 6 SAP implementations; reduced rework by 25% through structured fit-gap analysis.",
    "Delivered end-user training and post-go-live hypercare for 500+ users across manufacturing and logistics functions.",
]

EXP_GENPACT = [
    "Managed period-end close activities (GL, AP, AR) for 3 global clients; reduced close cycle errors by 30% through process standardisation and reconciliation automation.",
    "Resolved 200+ SAP FICO tickets (P1–P3) with 98% SLA compliance; recognised as top performer for 2 consecutive years.",
]

# ── DOCX FORMATTING HELPERS ──────────────────────────────────────────────────

def _set_font(run, name="Calibri", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def _add_section_heading(doc, text):
    """Add a styled section heading — ATS-safe (standard Heading style)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, size=11, bold=True, color=(31, 73, 125))
    # Underline via bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def _add_bullet(doc, text, indent=0.25):
    """Add an ATS-safe bullet point as a plain paragraph with • prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-indent)
    run = p.add_run(f"\u2022  {text}")
    _set_font(run, size=9.5)
    return p

def _add_normal(doc, text, bold=False, size=10, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p

def _add_job_header(doc, title, company, dates, location):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"{title}  |  ")
    _set_font(r1, size=10, bold=True)
    r2 = p.add_run(f"{company}  |  {dates}  |  {location}")
    _set_font(r2, size=10, bold=False)
    return p

# ── MAIN BUILDER ─────────────────────────────────────────────────────────────

def generate(post: dict) -> str:
    """
    Generate a tailored ATS resume for one recruiter post.
    Returns path to the generated .docx file.
    """
    jd_text  = post.get("full_post_text") or post.get("post_text") or ""
    company  = post.get("company") or ""
    email    = (post.get("email") or "noemail").replace("@", "_at_").replace(".", "_")
    keyword  = ""
    sk = post.get("search_keyword")
    if isinstance(sk, dict):
        keyword = sk.get("search", "")
    elif isinstance(sk, str):
        keyword = sk

    scores   = detect_roles(jd_text)
    summary  = pick_summary(scores)
    achievements = rank_achievements(scores)
    jd_kws   = extract_jd_keywords(jd_text)
    location_pref = detect_location_preference(jd_text)

    # Build title line — inject role from keyword if available
    title = TITLE_BASE
    if keyword:
        role_map = {
            "SAP S4 HANA Project Manager": "SAP S/4HANA Program Manager  |  Data Migration Specialist  |  15+ Years",
            "SAP Project Manager":         "SAP Program Manager  |  S/4HANA & ECC Expert  |  15+ Years",
            "SAP SD MM Consultant":        "SAP SD / MM Functional Consultant  |  S/4HANA Specialist  |  15+ Years",
            "SAP Data Migration Consultant":"SAP Data Migration Lead  |  LSMW / LTMC / SLT / CPI  |  15+ Years",
        }
        title = role_map.get(keyword, TITLE_BASE)

    # Inject JD keywords into skills section
    extra_skills = ""
    if jd_kws:
        # Only add ones not already in skills
        existing = " ".join(SKILLS_SECTIONS.values()).lower()
        new_kws = [k for k in jd_kws if k.lower() not in existing]
        if new_kws:
            extra_skills = " | ".join(new_kws)

    # ── CREATE DOCUMENT ──────────────────────────────────────────────────────
    doc = Document()

    # Page margins — narrow for space efficiency
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Set default font for entire document
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── HEADER ───────────────────────────────────────────────────────────────
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after  = Pt(2)
    r = p_name.add_run(NAME)
    _set_font(r, size=18, bold=True, color=(31, 73, 125))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(2)
    r = p_title.add_run(title)
    _set_font(r, size=11, bold=False)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after  = Pt(2)
    r = p_contact.add_run(CONTACT)
    _set_font(r, size=9.5)

    p_avail = doc.add_paragraph()
    p_avail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_avail.paragraph_format.space_before = Pt(0)
    p_avail.paragraph_format.space_after  = Pt(4)
    r = p_avail.add_run(AVAILABILITY)
    _set_font(r, size=9.5, bold=True)

    # ── CAREER SUMMARY ───────────────────────────────────────────────────────
    _add_section_heading(doc, "Career Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(summary)
    _set_font(r, size=9.5)

    # ── KEY ACHIEVEMENTS ─────────────────────────────────────────────────────
    _add_section_heading(doc, "Key Achievements")
    for ach in achievements:
        _add_bullet(doc, ach)

    # ── CORE COMPETENCIES & SKILLS ───────────────────────────────────────────
    _add_section_heading(doc, "Core Competencies & Technical Skills")
    for category, skills_text in SKILLS_SECTIONS.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(f"{category}: ")
        _set_font(r1, size=9.5, bold=True)
        r2 = p.add_run(skills_text)
        _set_font(r2, size=9.5)

    # Inject JD-specific extra keywords if any
    if extra_skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run("JD-Matched Skills: ")
        _set_font(r1, size=9.5, bold=True)
        r2 = p.add_run(extra_skills)
        _set_font(r2, size=9.5)

    # ── PROFESSIONAL EXPERIENCE ──────────────────────────────────────────────
    _add_section_heading(doc, "Professional Experience")

    # Autodesk
    _add_job_header(doc,
        "SAP S/4HANA Program Manager",
        "Autodesk Inc.",
        "May 2015 – Present",
        "Remote, India"
    )
    # Reorder Autodesk bullets based on role
    autodesk_bullets = list(EXP_AUTODESK)
    if scores.get("data_migration", 0) >= scores.get("program_manager", 0):
        # Move data migration bullet to top
        dm_bullet = autodesk_bullets[0]  # already first
        autodesk_bullets = sorted(
            autodesk_bullets,
            key=lambda b: (
                -2 if any(kw in b.lower() for kw in ["migration", "ltmc", "slt", "cpi", "lsmw", "etl"]) else
                -1 if any(kw in b.lower() for kw in ["cutover", "reconciliation", "mdg"]) else
                0
            )
        )
    for bullet in autodesk_bullets:
        _add_bullet(doc, bullet)

    # Lava
    _add_job_header(doc,
        "SAP Manager – FICO & Data Migration",
        "Lava International Ltd.",
        "2013 – 2015",
        "Remote, India"
    )
    for bullet in EXP_LAVA:
        _add_bullet(doc, bullet)

    # Sony
    _add_job_header(doc,
        "SAP Program Lead – FICO",
        "Sony India Pvt. Ltd.",
        "2011 – 2013",
        "Remote, India"
    )
    for bullet in EXP_SONY:
        _add_bullet(doc, bullet)

    # iGATE
    _add_job_header(doc,
        "SAP Implementation Consultant – FICO / MM / SD",
        "iGATE Patni | Semantic Space Technologies",
        "2008 – 2011",
        "Remote, India"
    )
    for bullet in EXP_IGATE:
        _add_bullet(doc, bullet)

    # Genpact
    _add_job_header(doc,
        "SAP FICO Analyst",
        "Genpact",
        "2006 – 2008",
        "Remote, India"
    )
    for bullet in EXP_GENPACT:
        _add_bullet(doc, bullet)

    # ── EDUCATION ────────────────────────────────────────────────────────────
    _add_section_heading(doc, "Education & Certifications")
    _add_bullet(doc, "MBA – Information Technology")
    _add_bullet(doc, "B.Com  |  B.Sc (Computer Science)  |  Diploma – Computer Engineering")

    # ── SAVE ─────────────────────────────────────────────────────────────────
    # Sanitize email for filename
    safe_email = re.sub(r"[^a-zA-Z0-9_]", "_", email)[:40]
    filename   = f"Harsh_Madan_SAP_{safe_email}.docx"
    filepath   = os.path.join(TEMP_DIR, filename)
    doc.save(filepath)
    return filepath


# ── CLI TEST ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    # Test with a sample post
    test_post = {
        "full_post_text": (
            "We are hiring SAP Data Migration Consultant with 10+ years experience. "
            "Must have hands-on LSMW, LTMC, SLT, SAP CPI. S/4HANA experience preferred. "
            "FICO module knowledge required. Cutover planning and reconciliation experience. "
            "Remote role. Contact: test@example.com"
        ),
        "company": "Test Company",
        "email": "test@example.com",
        "recruiter_name": "Test Recruiter",
        "search_keyword": {"search": "SAP Data Migration Consultant"},
    }
    path = generate(test_post)
    print(f"Generated: {path}")
    # Verify it opens
    doc = Document(path)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print("SUCCESS — resume generated and verified.")
