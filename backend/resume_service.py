import os
import json
from docx import Document
from openai_service import AIService
from uuid import uuid4

class ResumeService:
    def __init__(self, s3_client=None):
        self.ai_service = AIService()
        # S3 client would be injected here for commercial cloud storage
        self.s3_client = s3_client 

    def extract_text_from_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)

    def generate_tailored_resume(self, master_resume_path: str, job_description: str, output_dir: str):
        # 1. Extract text from the master docx
        resume_text = self.extract_text_from_docx(master_resume_path)

        # 2. Get AI tailoring instructions (JSON)
        ai_response_json = self.ai_service.analyze_and_tailor(resume_text, job_description)
        
        try:
            tailoring_data = json.loads(ai_response_json)
        except Exception as e:
            print(f"Failed to parse AI response: {str(e)}")
            return None

        # 3. Create a tailored version of the docx
        # For the commercial version, we'll clone the master and replace the summary.
        # This preserves all other formatting, headers, and footers.
        doc = Document(master_resume_path)
        
        # Simple replacement logic for the MVP
        # In a real enterprise app, we'd search for a specific heading or placeholder
        summary_replaced = False
        for para in doc.paragraphs:
            # Look for common summary headers
            if any(h in para.text.upper() for h in ["SUMMARY", "PROFILE", "PROFESSIONAL SUMMARY"]):
                # The next paragraph (or the rest of this one) is the summary
                summary_replaced = True
                continue
            
            if summary_replaced:
                para.text = tailoring_data.get("new_summary", para.text)
                summary_replaced = False
                break
        
        # 4. Save the tailored file
        new_filename = f"tailored_resume_{uuid4().hex}.docx"
        output_path = os.path.join(output_dir, new_filename)
        doc.save(output_path)
        
        return {
            "file_path": output_path,
            "match_score": tailoring_data.get("match_score", 0),
            "tailoring_data": tailoring_data
        }
