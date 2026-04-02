"""
Daily SAP Job Application - Outlook Draft Creator
Generates tailored resumes + Outlook drafts for today's LinkedIn leads.
Run via Create_SAP_Job_Drafts.bat in the Linkdin Job Application folder.
"""

import win32com.client as win32
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FOLDER    = r'C:\Users\madan\OneDrive\Desktop\Linkdin Job Application'
SIGNATURE = (
    "\n\nKind regards,\n"
    "Harsh Madan\n"
    "+91 96679 64756\n"
    "Madan.harsh1984@gmail.com\n"
    "LinkedIn: https://sg.linkedin.com/in/harsh-madan-b818113b"
)

# ─────────────────────────────────────────
# RESUME BUILDER
# ─────────────────────────────────────────
def add_heading(doc, text, size=14, bold=True, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold     = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    # underline separator
    doc.add_paragraph('─' * 80)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text).font.size = Pt(10)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10)

def build_resume(filename, summary, extra_skills, role_focus_bullets):
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin   = Pt(54)
        section.right_margin  = Pt(54)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run('HARSH MADAN')
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('SAP S/4HANA Program Manager | Data Migration Specialist | ERP Transformation Leader').font.size = Pt(10)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('+91 96679 64756  |  Madan.harsh1984@gmail.com  |  https://sg.linkedin.com/in/harsh-madan-b818113b').font.size = Pt(9)

    doc.add_paragraph()

    # Professional Summary
    add_section_title(doc, 'Professional Summary')
    add_body(doc, summary)

    # Core Competencies
    add_section_title(doc, 'Core Competencies')
    competencies = [
        'SAP S/4HANA Program & Project Management',
        'Full Lifecycle ERP Implementation (Plan to Go-Live)',
        'SAP Data Migration (LSMW | LTMC | SLT)',
        'Cutover Planning & Go-Live Management',
        'SAP Activate Methodology',
        'Procure-to-Pay (P2P) | Order-to-Cash (O2C)',
        'SIT / UAT / Business Sign-Off',
        'Stakeholder & Vendor Management',
        'Risk Management & Issue Resolution',
        'Agile / Hybrid Delivery Models',
        'SAP Modules: FICO | MM | SD | MDG | SAP CPI',
        'Tools: SAP Solution Manager | JIRA | Smartsheet | ServiceNow',
    ] + extra_skills
    for c in competencies:
        add_bullet(doc, c)

    # Professional Experience
    add_section_title(doc, 'Professional Experience')

    roles = [
        {
            'title':   'SAP S/4HANA Program Manager',
            'company': 'Autodesk',
            'dates':   'May 2015 – Present',
            'bullets': role_focus_bullets + [
                'Orchestrated end-to-end SAP S/4HANA ERP implementation including planning, execution, cutover, and stabilization phases',
                'Led P2P data migration covering Vendor Master, Purchase Orders, GR/IR reconciliation, and invoice processing integrated with SAP FI/AP',
                'Managed multi-disciplinary global project teams; coordinated with business stakeholders, technical teams, and external vendors',
                'Executed ETL processes using LSMW, LTMC, and SLT; achieved zero data inconsistency at go-live across multiple deployments',
                'Drove SIT/UAT cycles, managed defect resolution, and obtained business sign-off for all critical data objects',
                'Applied SAP Activate methodology; maintained project governance, risk register, and change control processes',
                'Managed project budgets, timelines, and scope in complex multi-country environments',
            ]
        },
        {
            'title':   'SAP FICO Consultant',
            'company': 'Sony India',
            'dates':   'Sep 2014 – May 2015',
            'bullets': [
                'Led GL, AP, AR, and Asset Accounting data migration from legacy ECC to S/4HANA',
                'Performed data validation, reconciliation, and obtained business sign-off on all financial data objects',
            ]
        },
        {
            'title':   'SAP Manager',
            'company': 'Lava International',
            'dates':   'Jan 2014 – Sep 2014',
            'bullets': [
                'Managed SAP implementation and data migration planning; optimized P2P processes',
                'Designed data mapping from legacy systems to SAP S/4HANA',
            ]
        },
        {
            'title':   'SAP Program Lead',
            'company': 'Semantic Space Technologies',
            'dates':   'May 2011 – Jan 2014',
            'bullets': [
                'Led global SAP rollouts with multi-country data migration using LSMW, LTMC, and SLT',
            ]
        },
        {
            'title':   'SAP Implementation Consultant',
            'company': 'iGATE Patni',
            'dates':   'Apr 2009 – May 2011',
            'bullets': [
                'Delivered SAP ECC implementations; managed ETL pipelines and data verification',
            ]
        },
        {
            'title':   'SAP FICO Analyst',
            'company': 'Genpact',
            'dates':   'Aug 2007 – Mar 2009',
            'bullets': [
                'Financial data validation, reconciliation, and SAP reporting across FICO modules',
            ]
        },
    ]

    for role in roles:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r1 = p.add_run(role['title'] + '  |  ' + role['company'])
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run('   ' + role['dates'])
        r2.italic = True
        r2.font.size = Pt(9)
        for b in role['bullets']:
            add_bullet(doc, b)

    # Education
    add_section_title(doc, 'Education')
    for edu in [
        'MBA (Information Technology)',
        'B.Com (Bachelor of Commerce)',
        'B.Sc (Computer Science)',
        'Diploma in Computer Engineering',
    ]:
        add_bullet(doc, edu)

    path = os.path.join(FOLDER, filename)
    doc.save(path)
    print('  Resume saved: ' + filename)
    return path


# ─────────────────────────────────────────
# TODAY'S LEADS
# ─────────────────────────────────────────
jobs = [
    {
        # REAL JOB — Rahul Shastri / Vera Zinc
        'to':       'rahul@verazinc.com',
        'subject':  'Application \u2013 Senior SAP Project Manager | Atlanta, GA | Harsh Madan | 15+ Years SAP S/4HANA',
        'resume':   'Harsh_Madan_SAP_PM_VeraZinc.docx',
        'summary': (
            'Results-driven SAP S/4HANA Program Manager with 15+ years of experience leading full lifecycle SAP ERP '
            'implementations from planning through cutover and stabilization. Currently serving as SAP S/4HANA Program '
            'Manager at Autodesk, with over 10 years of dedicated SAP ERP PM experience including the last several years '
            'exclusively on SAP S/4HANA implementations. Proven expertise in managing large multi-disciplinary teams, '
            'coordinating with C-suite stakeholders and external vendors, driving project governance, risk management, '
            'and budget control in complex enterprise environments.'
        ),
        'extra_skills': [
            'Project Governance & Risk Management',
            'Budget & Scope Management',
            'C-Suite Stakeholder Communication',
            'Vendor Coordination',
        ],
        'focus_bullets': [
            'Served as Senior SAP Project Manager for 10+ years, with last 5 years dedicated to full-lifecycle SAP S/4HANA implementations',
            'Led full cycle implementations including planning, execution, cutover, and post go-live stabilization',
            'Coordinated with C-suite and business stakeholders to align SAP delivery with strategic objectives',
            'Managed project budgets, timelines, and scope across complex multi-vendor, multi-country environments',
            'Established project governance frameworks, risk registers, and change control processes',
        ],
        'body': (
            'Dear Rahul,\n\n'
            'I came across your LinkedIn post for the Senior SAP Project Manager role in Atlanta, GA (12+ months) '
            'and would like to submit my profile for your consideration.\n\n'
            'I bring 15+ years of SAP experience with 10+ years in dedicated SAP ERP Project/Program Management, '
            'including the last several years leading full-lifecycle SAP S/4HANA implementations. In my current role '
            'as SAP S/4HANA Program Manager at Autodesk, I have led end-to-end ERP deliveries covering planning, '
            'execution, cutover, and stabilization \u2014 managing large multi-disciplinary teams, coordinating with '
            'C-suite stakeholders, and overseeing external vendors while maintaining strict governance over budgets, '
            'timelines, and scope.\n\n'
            'My expertise spans SAP S/4HANA data migration (LSMW/LTMC/SLT), integration management, risk resolution, '
            'and change control \u2014 consistently delivering go-lives with zero data inconsistency. I am well-versed in '
            'translating complex technical concepts into business-friendly language for executive audiences, which I '
            'understand is a key requirement for this role.\n\n'
            'Please find my tailored resume attached. I am happy to schedule a call at your earliest convenience to '
            'discuss alignment with your client\'s requirements.'
        ),
    },
    {
        # Koya Consult — Sammy David (bench sales recruiter)
        'to':       'Sammy@koyaconsult.com',
        'subject':  'SAP S/4HANA Consultant Available \u2013 Harsh Madan | 15+ Years | FICO | MM | Program Manager',
        'resume':   'Harsh_Madan_SAP_Koya_Sammy.docx',
        'summary': (
            'Seasoned SAP S/4HANA Program Manager and Consultant with 15+ years of experience across FICO, MM, SD, '
            'and data migration. Available for C2C / contract engagements. Expertise in full lifecycle S/4HANA '
            'implementations, cutover management, data migration (LSMW/LTMC/SLT), SIT/UAT, and global team leadership. '
            'Currently at Autodesk as SAP S/4HANA Program Manager. Immediate availability for the right opportunity.'
        ),
        'extra_skills': ['C2C / Contract Engagement', 'Immediate Availability'],
        'focus_bullets': [
            'Available for C2C / contract SAP S/4HANA Program Manager and Consultant roles',
            'Deep expertise in SAP FICO, MM, SD modules with 15+ full lifecycle implementation experience',
        ],
        'body': (
            'Dear Sammy,\n\n'
            'I noticed your LinkedIn post at Koya Consult regarding available SAP FICO / ABAP / S4 HANA consultants '
            'and wanted to reach out as a highly experienced SAP professional open to contract opportunities.\n\n'
            'I am an SAP S/4HANA Program Manager with 15+ years of experience across FICO, MM, SD, and large-scale '
            'ERP implementations. My background includes full lifecycle S/4HANA deliveries, cutover management, data '
            'migration (LSMW/LTMC/SLT), SIT/UAT coordination, and global team leadership \u2014 currently at Autodesk.\n\n'
            'I am open to C2C contract engagements for SAP S/4HANA Program Manager, SAP FICO, or SAP MM roles. '
            'Please find my resume attached \u2014 I would be happy to connect and discuss any requirements that fit my profile.'
        ),
    },
    {
        # Koya Consult — Parker / Durga Prasad
        'to':       'parker@koyaconsult.com',
        'subject':  'SAP S/4HANA Consultant Profile \u2013 Harsh Madan | 15+ Years | Program Manager | FICO | MM',
        'resume':   'Harsh_Madan_SAP_Koya_Parker.docx',
        'summary': (
            'Experienced SAP S/4HANA Program Manager with 15+ years spanning FICO, MM, SD, and data migration. '
            'Available for contract / C2C engagements. Proven track record in full lifecycle S/4HANA implementations, '
            'P2P transformation, cutover planning, and multi-country rollouts. Currently at Autodesk. '
            'Open to Technical Project Manager, Program Manager, and SAP Consultant roles.'
        ),
        'extra_skills': ['C2C / Contract Engagement', 'Technical Program Manager'],
        'focus_bullets': [
            'Open to Technical Project Manager and SAP S/4HANA Consultant contract roles',
            '15+ years of FICO, MM, S/4HANA delivery across multiple global organizations',
        ],
        'body': (
            'Dear Parker,\n\n'
            'I came across the Koya Consulting LinkedIn post about available SAP and IT consultants and wanted to '
            'add my profile to your distribution list for C2C / contract opportunities.\n\n'
            'I am an SAP S/4HANA Program Manager with 15+ years of experience across FICO, MM, SD, and large-scale '
            'ERP transformations. My expertise includes full lifecycle S/4HANA implementations, P2P data migration, '
            'cutover management, and global multi-country rollouts \u2014 currently delivering as Program Manager at Autodesk.\n\n'
            'I am available for Technical Project Manager, SAP FICO / MM Consultant, and SAP S/4HANA Program Manager '
            'contract roles. Please find my resume attached and do add me to your distribution list for relevant openings.'
        ),
    },
    {
        # Koya Consult — James / Rupas
        'to':       'james@koyaconsult.com',
        'subject':  'SAP S/4HANA Program Manager Available \u2013 Harsh Madan | 15+ Years | FICO | MM | C2C',
        'resume':   'Harsh_Madan_SAP_Koya_James.docx',
        'summary': (
            'Highly experienced SAP S/4HANA Program Manager with 15+ years in FICO, MM, SD, data migration, and '
            'global ERP rollouts. Available for C2C / contract engagements in Program Manager, Technical Project Manager, '
            'or SAP FICO / MM Consultant roles. Currently at Autodesk with a strong record of on-time, zero-defect go-lives.'
        ),
        'extra_skills': ['C2C / Contract Engagement', 'Scrum Master exposure', 'Program Manager'],
        'focus_bullets': [
            'Available for C2C contract as SAP S/4HANA Program Manager or SAP FICO / MM Consultant',
            '15+ years of ERP delivery at global organizations including Autodesk, Sony, iGATE Patni',
        ],
        'body': (
            'Dear Rupas,\n\n'
            'I came across your LinkedIn hotlist post at Koya Consult and wanted to share my profile for any '
            'relevant SAP S/4HANA / FICO / Program Manager contract requirements in your pipeline.\n\n'
            'I am a seasoned SAP S/4HANA Program Manager with 15+ years of experience covering SAP FICO, MM, SD, '
            'data migration, cutover planning, and full lifecycle ERP implementations. Currently at Autodesk, I have '
            'consistently delivered S/4HANA go-lives with zero data inconsistency, managing cross-functional global '
            'teams and C-suite stakeholders.\n\n'
            'I am open to C2C / contract roles as Technical Program Manager, SAP FICO Consultant, or SAP S/4HANA '
            'Consultant. Please find my resume attached and feel free to reach out to discuss any active requirements.'
        ),
    },
]


# ─────────────────────────────────────────
# BUILD RESUMES + OUTLOOK DRAFTS
# ─────────────────────────────────────────
print('=' * 60)
print('  Daily SAP Job Application — ' + __import__('datetime').date.today().strftime('%d %b %Y'))
print('=' * 60)
print(f'\nBuilding {len(jobs)} resumes and Outlook drafts...\n')

# Connect Outlook
try:
    outlook = win32.GetActiveObject('Outlook.Application')
    print('Outlook: connected to running instance.\n')
except Exception:
    outlook = win32.Dispatch('Outlook.Application')
    print('Outlook: launched.\n')

success = 0
for i, job in enumerate(jobs, 1):
    print(f'[{i}/{len(jobs)}] {job["to"]}')
    try:
        # Build resume
        att_path = build_resume(job['resume'], job['summary'], job['extra_skills'], job['focus_bullets'])

        # Create Outlook draft
        mail          = outlook.CreateItem(0)
        mail.To       = job['to']
        mail.Subject  = job['subject']
        mail.Body     = job['body'] + SIGNATURE
        mail.Attachments.Add(att_path)
        mail.Save()
        print('  Draft saved to Outlook Drafts.\n')
        success += 1
    except Exception as e:
        print(f'  ERROR: {e}\n')

print('=' * 60)
print(f'DONE: {success}/{len(jobs)} drafts saved to Outlook Drafts folder.')
print('Open Outlook > Drafts to review and send.')
print('=' * 60)
input('\nPress Enter to close...')
